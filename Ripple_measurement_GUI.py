import io
from logging import raiseExceptions
import sys
import docx
import paramiko
import datetime
import pandas as pd
from tkinter import *
import pyvisa as visa
from docx.shared import Cm
from tkinter import filedialog

def connect():
    # Prisijungimas prie prietaisų: Elektroninės apkrovos ir Osciloskopo
    global rm, device, second_device, scope, e_load
    rm = visa.ResourceManager()
    try:
        #Adreso išskaidymas ir patikrinimas, kad būtų įrašomi tik skaitmenys
        numbers=first_address_entry.get()
        first_number=numbers[0:3]
        second_number=numbers[4:7]
        third_number=numbers[8:10]
        fourth_number=numbers[11:13]
        try:
            int(first_number)
            int(second_number)
            int(third_number)
            int(fourth_number)
        except ValueError:
            info('Address contains only numbers')
        device = rm.open_resource(f"TCPIP::{first_address_entry.get()}::INSTR")#192.168.91.20  
    except visa.VisaIOError:
        info("Electronic Load is not connected")
        device= None

    try: 
        #Adreso išskaidymas ir patikrinimas, kad būtų įrašomi tik skaitmenys
        numbers_2=second_address_entry.get()
        first_number_2=numbers_2[0:3]
        second_number_2=numbers_2[4:7]
        third_number_2=numbers_2[8:10]
        fourth_number_2=numbers_2[11:13]
        try:
            int(first_number_2)
            int(second_number_2)
            int(third_number_2)
            int(fourth_number_2)
        except ValueError:
            info('Address contains only numbers')
        
        second_device = rm.open_resource(f"TCPIP::{second_address_entry.get()}::INSTR")#192.168.91.25
    except visa.VisaIOError:
        info("Oscilloscope is not connected")  
        second_device=None 
    # Prietaisams prisikiriami kintamieji Osciloskopui - scope, E-apkrovai - e_load 
    if device== None and second_device==None:
        info("Devices not connected")
    if device != None:
        split_word_eload=device.query('*IDN?')
        e_load = None
        Name = split_word_eload.split()
        if Name[0] == "Siglent":
            e_load=device
            info('Electronic Load - Connected')
            info(e_load.query('*IDN?'))
    if second_device != None:
        split_word_scope=second_device.query('*IDN?')
        scope = None
        Name2 = split_word_scope.split()
        if Name2[0] == "RIGOL": 
            scope=second_device
            info('Oscilloscope - Connected')
            info(scope.query('*IDN?'))

def connect_trb():
    # Prisijungimas prie tinklo sietuvo
    global client
    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    #client.connect('192.168.91.30',22,'root','Kaunastlt1')
    try:
        #Adreso išskaidymas ir patikrinimas, kad būtų įrašomi tik skaitmenys
        numbers_3=third_address_entry.get()
        first_number_3=numbers_3[0:3]
        second_number_3=numbers_3[4:7]
        third_number_3=numbers_3[8:10]
        fourth_number_3=numbers_3[11:13]
        try:
            int(first_number_3)
            int(second_number_3)
            int(third_number_3)
            int(fourth_number_3)
        except ValueError:
            info('Address contains only numbers')
        client.connect(f"{third_address_entry.get()}",22,'root',f"{address_password_entry.get()}")
        info('Successful connection with TRB')

    except IOError :
        info('NOT Successful connection with TRB')
    except paramiko.ssh_exception.AuthenticationException:
        info('Wrong Password')

def close_connection():  
    # Atsijungimas nuo prietaisų
  try:
      if device!=None:
          rm.close()
          info('E_load and Scope resources closed')
      if second_device!=None:
          rm.close()
  except NameError:
      info('No connection with Scope or E_load')
      pass
  try:
        if client!=None:
            client.close()
            info('Gateway resource closed')
            
  except:
      info('No connection with TRB')
  
    
def enable_no_load():
    # No load mygtuko įgalinimas, prisijungus prie osciloskopo
    try:
        if second_device is None:
            pass
        else:no_load_checkbutton.configure(state='normal')
    except: NameError

def enable_load_measurements():
    # Full, dynamic ir long term mygtukų įgalinimas, prijungus apkrovą ir nustačius srovę
    try:
        if client!=None:
            for child in without_current[5:]:
                child.configure(state='normal')
    except NameError:
        pass
    try:
        if device!=None:
            full_load_checkbutton.configure(state='normal')
            dynamic_load_checkbutton.configure(state='normal')
            long_term_checkbutton.configure(state='normal')
            for child in current_entry_frame.winfo_children():
                child.configure(state='normal')
    except NameError:
        pass

def set_voltage():
    pass
def oscilloscope_parameters():
    # Standartiniai Osciloskopo matavimo parametrai
    global device
    try:
        if device != None:
            scope.write(':CHANnel1:VERNier OFF')
            scope.write(':TIMebase:VERNier OFF')
            scope.write(':CHANnel1:COUPling AC') 
            scope.write(':CHANnel1:BWLimit 20M')
            root.after(3000)
            scope.write(':SYSTem:KEY:PRESs TLEVel')
            scope.write(':SYSTem:KEY:PRESs VOFFset1')
    except: NameError

def multiplexer(channel_nr,select):
    # Kanalo dekodavimas
    global CH_nr
    #decimal value of select
    decode= (select[0] * 2) + (select[1] * 1)
    # getting the output of decimal value from inputs 
    CH_nr = channel_nr[decode]
    #print the output
    #channel_nr=[1,3,2,4]
    print(CH_nr)    
    pick_channel(CH_nr)    
    #multiplexer([1,3,2,4],[1,1])

def pick_channel(CH_nr):
    # Srovės įvedimo laukelio priskyrimas srovės kintamąjam "current", priklausomai nuo kanalo
    global current
    if CH_nr==1:

        current=current_1_entry.get()
        info(f"CH 1 current set: {current}")
        try:
            # Tikrinimas ar įvesta srovė, jei pasirinktas full, dynamic arba long term matavimas
            if full_load_variable.get() == 1 or dynamic_load_variable.get() == 1 or long_term_variable.get() == 1:
                if len(current_1_entry.get()) == 0:
                    info('Set Current CH1!')
                    raise Exception("Set Current CH1!")
                else:
                    current_1_entry.configure(state='disabled')
                    pass
        except NameError:
            info('Set current CH1')
    
    if CH_nr==2:
   
        current=current_2_entry.get()
        info(f"CH 2 current set: {current}")
        try:
            # Tikrinimas ar įvesta srovė, jei pasirinktas full, dynamic arba long term matavimas
            if full_load_variable.get() == 1 or dynamic_load_variable.get() == 1 or long_term_variable.get() == 1:
                if len(current_2_entry.get()) == 0:
                    info('Set Current CH2!')
                    raise Exception("Set Current CH2!")
                else:
                    current_2_entry.configure(state='disabled')
                    pass
        except NameError:
            info('Set current CH2')

    if CH_nr==3: 
   
        current=current_3_entry.get()
        
        info(f"CH 3 current set: {current}")
        try:
            # Tikrinimas ar įvesta srovė, jei pasirinktas full, dynamic arba long term matavimas
            if full_load_variable.get() == 1 or dynamic_load_variable.get() == 1 or long_term_variable.get() == 1:
                if len(current_3_entry.get()) == 0:
                    info('Set Current CH3!')
                    raise Exception("Set Current CH3!")
                else:
                    current_3_entry.configure(state='disabled')
                    pass
        except NameError:
            info('Set current CH3')
    if CH_nr==4:
      
        current=current_4_entry.get()
        info(f"CH 4 current set: {current}")
        try:
            # Tikrinimas ar įvesta srovė, jei pasirinktas full, dynamic arba long term matavimas
            if full_load_variable.get() == 1 or dynamic_load_variable.get() == 1 or long_term_variable.get() == 1:
                if len(current_4_entry.get()) == 0:
                    info('Set Current CH4!')
                    raise Exception("Set Current CH4!")
                else:
                    current_4_entry.configure(state='disabled')
                    pass
        except NameError:
            info('Set current CH4')
            
    set_current(current)
    return current
    
def set_current(current):
    # Įvestos srovės reikšmės nustatymas ir statiniame ir dinaminiame režime
    print ("current set", current) 
    try:
        e_load.write(f":SOURce:CURRent:LEVel:IMMediate {current}")
        e_load.write(f":SOURce:CURRent:TRANsient:BLEVel {current}")
    except NameError:
        print('Stringa')

def first_time():
    # Patikrinimas ar buvo pakeisti dinaminiai parametrai
    global dynamic_changed
    dynamic_changed=1
    Submit_btn.configure(state='disabled') 
 
def default_dynamic_parameters():
    # Dinaminių parametrų keitimui skitų laukų vartotojo sąsajoje sukūrimas
    global A_lvl_entry, A_wid_entry, B_wid_entry, Ris_slew_entry, Fal_slew_entry, current, Submit_btn
    Set_dyn_lbl=Label(Set_dyn_frame,text='Set DYN:')
    Set_dyn_lbl.grid(row=0,column=0)

    A_lvl_lbl=Label(Set_dyn_frame,text='A_lvl')
    A_lvl_lbl.grid(row=2,column=0,sticky=E)
    A_wid_lbl=Label(Set_dyn_frame,text='A_width[s]')
    B_wid_lbl=Label(Set_dyn_frame,text='B_width[s]')
    A_wid_lbl.grid(row=4,column=0,sticky=E)
    B_wid_lbl.grid(row=5,column=0,sticky=E)
    Ris_slew_lbl=Label(Set_dyn_frame,text='Ris_slew[A/us]')
    Fal_slew_lbl=Label(Set_dyn_frame,text='Fal_slew[A/us]')
    Ris_slew_lbl.grid(row=6,column=0,sticky=E)
    Fal_slew_lbl.grid(row=7,column=0,sticky=E)

    A_lvl_entry=Entry(Set_dyn_frame, width=12)
    # Pradinių (default) reikšmių įrašymas
    A_lvl_entry.insert(END, '0.03')
    A_lvl_entry.grid(row=2,column=1,sticky=W)
    A_wid_entry=Entry(Set_dyn_frame, width=12)
    B_wid_entry=Entry(Set_dyn_frame, width=12)
    # Pradinių (default) reikšmių įrašymas
    A_wid_entry.insert(END, '0.1')
    B_wid_entry.insert(END, '0.1')
    A_wid_entry.grid(row=4,column=1,sticky=W)
    B_wid_entry.grid(row=5,column=1,sticky=W)
    Ris_slew_entry=Entry(Set_dyn_frame, width=12)
    Fal_slew_entry=Entry(Set_dyn_frame, width=12)
    # Pradinių (default) reikšmių įrašymas
    Ris_slew_entry.insert(END, '2.5')
    Fal_slew_entry.insert(END, '2.5')
    Ris_slew_entry.grid(row=6,column=1,sticky=W)
    Fal_slew_entry.grid(row=7,column=1,sticky=W)

    Submit_btn=Button(Set_dyn_frame,text='Submit', command=lambda:[set_dynamic_parameters(),first_time()])
    Submit_btn.grid(row=8,column=1)

def set_dynamic_parameters():
    # Dinaminių parametrų nustatymas
    global A_wid_entry, B_wid_entry, half_period, second_half_period
    try:
        if device != None:
            try:
                if A_lvl_entry.get()>=str(0) and A_lvl_entry.get()<str(30.0):
                    if A_lvl_entry.get()==str(0):
                        e_load.write(f":SOURce:CURRent:TRANsient:ALEVel 0.0")
                    else:
                        e_load.write(f":SOURce:CURRent:TRANsient:ALEVel {A_lvl_entry.get()}")
                else:
                    info('Entry accept only numbers with point')
                    print('a_lvl else')
                if A_wid_entry.get()>str(0.0002) and A_wid_entry.get()<str(999.0):
                    half_period=float(A_wid_entry.get())
                    e_load.write(f":SOURce:CURRent:TRANsient:AWIDth {A_wid_entry.get()}")
                else:
                    info('Entry accept only numbers with point')
                    print('a_wid else')
                if B_wid_entry.get()>str(0.0002) and B_wid_entry.get()<str(999.0):
                    second_half_period=float(B_wid_entry.get())
                    e_load.write(f":SOURce:CURRent:TRANsient:BWIDth {B_wid_entry.get()}")
                else:
                    info('Entry accept only numbers with point')
                    print('b_wid else')
                if Ris_slew_entry.get()>str(0.001) and Ris_slew_entry.get()<str(2.5001):
                    e_load.write(f":SOURce:CURRent:TRANsient:SLEW:POSitive {Ris_slew_entry.get()}")
                else:
                    info('Entry accept only numbers with point')
                    print('ris else')
                if Fal_slew_entry.get()>str(0.001) and Fal_slew_entry.get()<str(2.5001):
                    e_load.write(f":SOURce:CURRent:TRANsient:SLEW:NEGative {Fal_slew_entry.get()}")
                else:
                    info('Entry accept only numbers with point')
                    print('fal else')
                e_load.write(':SOURce:FUNCtion:TRANsient CURRent')
                print ("Nustatome dinaminę apkrovą")
            except ValueError:
                info('Entry accept only numbers with point1')
                print('error else')

        else:
            info('Electronic Load is not connected')
    except NameError:
        info('Electronic Load is not connected') 

def auto():
    # Osciloskopo Autoscale funkcija 
    root.after(60)
    scope.write(':AUToscale')
    root.after(6000)   

def No_load():
    # Osciloskopo suderinimas matavimui be aprkovos
    print('No load')
    global Vrms,scope
    try:
        e_load.write(':SOURce:INPut:STATe OFF')
    except NameError:
        pass
    oscilloscope_parameters()
    auto()
    # Osciloskopo skalių parinkimas, kurosorių nustatymas į pradines pozicijas
    scope.write(f":CHANnel1:SCALe 0.05") 
    scope.write(':TIMebase:MAIN:SCALe 0.05')
    scope.write(':SYSTem:KEY:PRESs TLEVel')
    scope.write(':SYSTem:KEY:PRESs VOFFset1')
    # Skalių priderinimas pagal santykį ekrane
    for u in range(1,5):
        Vpp_1=float(scope.query(':MEASure:ITEM? VPP,CHANnel1'))
        vert_scale=float(scope.query(':CHANnel1:SCALe? '))
        Amplitude_Upper_limit=4.5*vert_scale #5*V_scale 3/1
        Amplitude_Lower_limit=2*vert_scale 
        root.after(3000)
        if Vpp_1>Amplitude_Upper_limit:
            scope.write(':SYSTem:KEY:DECRease VSCale1,1')
            print('Atsuka')
        if Vpp_1<Amplitude_Lower_limit:
            scope.write(':SYSTem:KEY:INCRease VSCale1,1')
            print('Suka')
        else: break
    half("no_load")

def Full_load():
    # Statinio režimo, osciloskopo skalių parinkimas, srovės įjungimas
    e_load.write(':SOURce:FUNCtion CURRent')
    e_load.write(':SOURce:INPut:STATe ON')
    print('Full load')
    root.after(1000)
    scope.write(':CHANnel1:COUPling AC')
    root.after(1000)
    scope.write(':SYSTem:KEY:PRESs VOFFset1')
    oscilloscope_parameters()
    root.after(2000)
    scope.write('RUN')
    scope.write(f":CHANnel1:SCALe 0.02") 
    scope.write(':TIMebase:MAIN:SCALe 0.0005')
    root.after(1000)
    scope.write(':SYSTem:KEY:PRESs TLEVel')
    scope.write(':SYSTem:KEY:PRESs VOFFset1')
    # Skalių priderinimas pagal santykį ekrane
    for u in range(1,5):
        Vpp_1=float(scope.query(':MEASure:ITEM? VPP,CHANnel1'))
        vert_scale=float(scope.query(':CHANnel1:SCALe? '))
        Amplitude_Upper_limit=5*vert_scale 
        Amplitude_Lower_limit=3*vert_scale 
        root.after(3000)
        if Vpp_1>Amplitude_Upper_limit:
            scope.write(':SYSTem:KEY:DECRease VSCale1,1')
            print('Atsuka')
        if Vpp_1<Amplitude_Lower_limit:
            scope.write(':SYSTem:KEY:INCRease VSCale1,1')
            print('Suka')
        else: break
    half("full_load")

def Dynamic_load():
    global A_wid_entry, B_wid_entry, dynamic_changed
    print('Dynamic load')
    # Dinaminio režimo, apkrovos reikšmių, osciloskopo skalių parinkimas, srovės įjungimas
    oscilloscope_parameters()
    try:
        if dynamic_changed==1:
            pass
        else: 
            default_dynamic_parameters()
            set_dynamic_parameters()
            print('pradines reiksmes')
    except NameError:
        dynamic_changed=0
        
    root.after(2000)
    e_load.write(':SOURce:INPut:STATe ON')
    scope.write('RUN')
    time_scale_dyn=round((half_period+second_half_period)/2,1)
    print(time_scale_dyn)
    scope.write(f":TIMebase:MAIN:SCALe {time_scale_dyn}")
    scope.write(f":CHANnel1:SCALe 0.5") 
    scope.write(':SYSTem:KEY:PRESs VOFFset1')
    root.after(2000)
    # Skalių priderinimas pagal santykį ekrane
    for u in range(1,7):
        Vpp_1=float(scope.query(':MEASure:ITEM? VPP,CHANnel1'))
        vert_scale=float(scope.query(':CHANnel1:SCALe? '))
        Amplitude_Upper_limit=5*vert_scale #5*V_scale 3/1
        Amplitude_Lower_limit=2*vert_scale 
        root.after(3000)
        if Vpp_1>Amplitude_Upper_limit:
            scope.write(':SYSTem:KEY:DECRease VSCale1,1')
            print('Atsuka')
        if Vpp_1<Amplitude_Lower_limit:
            scope.write(':SYSTem:KEY:INCRease VSCale1,1')
            print('Suka')
        else: 
            print('Vpp_1',Vpp_1)
            break
    root.after(1000)
    half("dynamic_load")
  
def Long_term():
    pass
    """
    print('Long term') 
    oscilloscope_parameters()
    # Statinio matavimo režimo nustatymas, skalių parinkimas, srovės įjungimas
    e_load.write(':SOURce:FUNCtion CURRent')
    scope.write('RUN')
    e_load.write(':SOURce:INPut:STATe ON')
    scope.write(f":CHANnel1:SCALe 0.02") 
    scope.write(':TIMebase:MAIN:SCALe 0.0005')
    scope.write(':SYSTem:KEY:PRESs TLEVel')
    scope.write(':SYSTem:KEY:PRESs VOFFset1')
    # Skalių priderinimas pagal santykį ekrane
    for u in range(1,5):
        Vpp_1=float(scope.query(':MEASure:ITEM? VPP,CHANnel1'))
        vert_scale=float(scope.query(':CHANnel1:SCALe? '))
        Amplitude_Upper_limit=5*vert_scale #5*V_scale 3/1
        Amplitude_Lower_limit=3*vert_scale 
        root.after(3000)
        if Vpp_1>Amplitude_Upper_limit:
            scope.write(':SYSTem:KEY:DECRease VSCale1,1')
            print('Atsuka')
        if Vpp_1<Amplitude_Lower_limit:
            scope.write(':SYSTem:KEY:INCRease VSCale1,1')
            print('Suka')
        else: break
    # Laikas kuriam praėjus bus atliekamas matavimas 1000=1s (1h)
    root.after(3600000)
    half("long_term")
    """

def cancel():
    # Srovės laukelių įvedimo ir submit mygtuko normalizavimas
    try:
        Submit_btn.configure(state='normal')
    except NameError:
        pass
    current_1_entry.configure(state='normal')
    current_2_entry.configure(state='normal')
    current_3_entry.configure(state='normal')
    current_4_entry.configure(state='normal')

def half(measurement):
    # Atliekami AC Vpp ir DC Vrms matavimai 
    try:
        print("Calibration is done, starting to measure")
        global Vpp_avg, Vrms, image_stream, vrms_no_load, vpp_no_load, vrms_full_load, vpp_full_load, vrms_dynamic_load, vpp_dynamic_load
        # Leidimas atlikti matavimus
        scope.write('RUN')
        root.after(3000)
        scope.write(':MEASure:ITEM VRMS,CHANnel1')
        scope.write(':MEASure:ITEM VPP,CHANnel1') 
        #root.after(3000)
        # Matavimai
        Vpp=float(scope.query(':MEASure:ITEM? VPP,CHANnel1'))
        if Vpp>1e+37:
            root.after(1000)
            del Vpp
            scope.write(':MEASure:ITEM VPP,CHANnel1')
            Vpp=float(scope.query(':MEASure:ITEM? VPP,CHANnel1'))
        #root.after(300)
        scope.write(':MEASure:ITEM VPP,CHANnel1') 
        Vpp2=float(scope.query(':MEASure:ITEM? VPP,CHANnel1'))
        if Vpp2>1e+37:
            root.after(1000)
            del Vpp2
            scope.write(':MEASure:ITEM VPP,CHANnel1') 
            Vpp2=float(scope.query(':MEASure:ITEM? VPP,CHANnel1'))
        root.after(200)
        scope.write(':MEASure:ITEM VPP,CHANnel1')
        Vpp3=float(scope.query(':MEASure:ITEM? VPP,CHANnel1'))
        root.after(200)
        scope.write(':MEASure:ITEM VPP,CHANnel1')
        Vpp4=float(scope.query(':MEASure:ITEM? VPP,CHANnel1'))
        root.after(200)
        scope.write(':MEASure:ITEM VPP,CHANnel1')
        Vpp5=float(scope.query(':MEASure:ITEM? VPP,CHANnel1'))
        
        scope.write(":DISP:DATA?")
        bmpdata = scope.read_raw()[2+9:]
        image_stream = io.BytesIO(bmpdata)

        scope.write(':CHANnel1:COUPling DC') 
        root.after(2000)
        auto()
        Vrms=float(scope.query(':MEASure:ITEM? VRMS,CHANnel1'))

        print(f"Vpp: {Vpp} V, DC Vrms: {Vrms} V")
        # Išmatuotų verčių priskyrimas GUI kintamiesiems
        Vpp_avg=round((Vpp+Vpp2+Vpp3+Vpp4+Vpp5)/5,4)
        print(f"Vpp_avg: {Vpp_avg} V, DC Vrms: {Vrms} V")

        if measurement == "no_load":
            vrms_no_load=round(Vrms,4)
            vpp_no_load=Vpp_avg
            print('saugojimui', vrms_no_load, vpp_no_load)
        if measurement == "full_load":
            vrms_full_load=Vrms
            vpp_full_load=Vpp_avg
            print('saugojimui', vrms_full_load, vpp_full_load)
        if measurement == "dynamic_load":
            vrms_dynamic_load=round(Vrms,4)
            vpp_dynamic_load=Vpp_avg
            print('saugojimui', vrms_dynamic_load, vpp_dynamic_load)
        if measurement == "long_term":
            vrms_long_term=round(Vrms,4)
            vpp_long_term=Vpp_avg
            print('saugojimui', vrms_long_term, vpp_long_term)

        scope.write(':SYSTem:KEY:PRESs MOFF')
        #root.after(2000)
        print("Done.")
    except:
        e_load.write(':SOURce:INPut:STATe OFF')


def info(message):
    # Funkcija žinutės išvedimui į teksto laukelį
    text_box.insert(END, '\n')
    text_box.insert('1.0', f"\n{message}\n")
    

def first():
    # Funkcija tikrinimui ar jau buvo sukurtas dokumentas
    first.has_been_called = True


def create_doc():
    # Naujo dokumento sukūrimas, užvadinimas pagal datą
    global flag,full_name, doc_name, doc
    doc = docx.Document()
    doc_name =datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    full_name=f"Ripple_{doc_name}.docx"
    doc.save(full_name)


def get_data(type):
    # Įrašo duomenis į dokumentą(Sukuria naują)
    global flag, image_stream, Vpp_avg
    first()
  
    if flag==1:
        pass
    else:
        print('saugom')
        save_file()
    
    data = {'Parameter': ['Ripple Vpp:','Vrms',],'Value': [Vpp_avg,Vrms]}
    df = pd.DataFrame(data, columns = ['Parameter', 'Value'])
    # save_file(create_doc) funkcijos sukurto dokumento atidarymas
    doc = docx.Document(full_name)
    doc.add_heading('Ripple Measurement', 0)
    doc.add_heading(f"Supply test board channel: {channel}", 3)

    # Pasirinkto matavimo pavadinimas dokumente
    if type=="No Load":
        doc.add_heading('No Load measurement:', 3)
    if type=="Full Load":
        doc.add_heading('Full Load measurement:', 3)
    if type=="Dynamic Load":
        doc.add_heading('Dynamic Load measurement:', 3)
    if type=="Long Term":
        doc.add_heading('Long Term measurement:', 3)
    else:
        pass
    doc.add_heading('Oscilloscope Image:', 3)
    # Osciloskopo vaizdo įrašymas į dokumentą
    doc.add_picture(image_stream, width=Cm(17), height=Cm(10))
    doc.add_heading('Measurement data:', 3)
    # Duomenų surašymas į dokumentą
    t = doc.add_table(df.shape[0]+1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])

    if not directory:
        print('No specific path')
        doc.save(full_name)
        return
    else:
        doc.save(full_name)
        path=directory+'/'+full_name
        doc.save(path)
          

def get_data_add(type1):
    # Įrašo (prirašo) duomenis į dokumentą (į jau sukurtą)
    root.after(2000)
    # Duomenų suformavimas į lentelę
    data = {'Parameter': ['Ripple Vpp:','Vrms',],'Value': [Vpp_avg,Vrms]}
    df = pd.DataFrame(data, columns = ['Parameter', 'Value'])
    # Get_data funkcijos sukurto dokumento atidarymas
    doc=docx.Document(full_name)
    # Pasirinkto matavimo pavadinimas dokumente
    if type1=="No Load":
        doc.add_heading('No Load measurement:', 3)
    if type1=="Full Load":
        doc.add_heading('Full Load measurement:', 3)
    if type1=="Dynamic Load":
        doc.add_heading('Dynamic Load measurement:', 3)
    if type1=="Long Term":
        doc.add_heading('Long Term measurement:', 3)
    else:
        pass

    doc.add_heading('Oscilloscope Image:', 3)
    root.after(1000)
    # Osciloskopo vaizdo įrašymas į dokumentą
    doc.add_picture(image_stream, width=Cm(17), height=Cm(10))
    doc.add_heading('Measurement data:', 3)
    # Duomenų surašymas į dokumentą
    t = doc.add_table(df.shape[0]+1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
    # Dokumento išsaugojimas
    doc.save(full_name)
    path=directory+'/'+full_name
    doc.save(path)
  

def table():
    # Atlikus no, full ir dynamic load matavimus suformuoja ir į dokumentą įrašo lentelę su duomenimis
    try:
        if no_load_variable.get()== 1 and full_load_variable.get()== 1 and dynamic_load_variable.get()== 1:
            try:
                if no_load_variable.get()== 1 and full_load_variable.get()== 1 and dynamic_load_variable.get()== 1:
                
                    if vpp_no_load!=str(0) and vrms_no_load!=str(0):
                        if vpp_full_load!=str(0) and vrms_full_load!=str(0):
                            if vpp_dynamic_load!=str(0) and vrms_dynamic_load!=str(0):
                                data = {'Sample': ['No load voltage, V:','Full load voltage, V', 'No load ripple peak to peak, V','Full load ripple peak to peak, V', 'Dynamic load ripple peak to peak',],'1': [vrms_no_load, vrms_full_load, vpp_no_load, vpp_full_load, vpp_dynamic_load]}
                                df = pd.DataFrame(data, columns = ['Sample', '1'])
                else: pass
            except NameError:
                pass
            doc=docx.Document(full_name)
            doc.add_heading('Measurement data:', 3)
            t = doc.add_table(df.shape[0]+1, df.shape[1])
            for j in range(df.shape[-1]):
                t.cell(0,j).text = df.columns[j]
            for i in range(df.shape[0]):
                for j in range(df.shape[-1]):
                    t.cell(i+1,j).text = str(df.values[i,j])
            # Dokumento išsaugojimas
            doc.save(full_name)
            path=directory+'/'+full_name
            doc.save(path)
    except:
        pass


def save_file():
    # Išsaugo sukurtą dokumentą, nurodytoje direktorijoje
    global path, directory, flag
    create_doc()
    try:
        # Klausia, kuriame aplanke išsaugoti dokumentą
        directory = filedialog.askdirectory()
        print('Direktorija', directory)
    except: PermissionError
     
    if not directory:
        print('Tuscia')
        doc.save(full_name)
        return
    else:
        path=directory+'/'+full_name
        flag=1
        print('Kelias', path)
        doc.save(path)
        #return flag


def clear_variables():
    # Ištrina kintamuoius naudojamus, kurti ir įrašyti į failą
    try:
        global full_name,doc_name,doc,directory,flag
        del full_name,doc_name,doc,directory,flag
    except NameError:
        pass

def measurements():
    # Dešifruoja kuriuos matavimus atlikti, tikrina ar iškviesti dokumento kūrimo funkcija ar prirašymo į dokumentą gale matavimo išjungia srovę
    first.has_been_called = False
    
    if no_load_variable.get() == 1:
        No_load()
        get_data("No Load")
        info('No load - Done')
    if full_load_variable.get() == 1:
        Full_load()
        if first.has_been_called == True:
            get_data_add("Full Load")
        else:
            get_data("Full Load") 
        e_load.write(':SOURce:INPut:STATe OFF')
        info('Full load - Done')
    if dynamic_load_variable.get() == 1:
        Dynamic_load()
        if first.has_been_called == True:
            get_data_add("Dynamic Load")
        else:
            get_data("Dynamic Load")
        e_load.write(':SOURce:INPut:STATe OFF')
        info('Dynamic load - Done')
    if long_term_variable.get()== 1:
        Long_term()
        if first.has_been_called == True:
            get_data_add("Long Term")
        else:
            get_data("Long Term")
        e_load.write(':SOURce:INPut:STATe OFF')
        info('Long term - Done')
    if no_load_variable.get()== 0 and full_load_variable.get()== 0 and dynamic_load_variable.get()== 0 and long_term_variable.get()== 0:
        info('No selected measurements')
    else:
        table()
        info('Measurements completed')
        first.has_been_called == False
        e_load.write(':SOURce:INPut:STATe OFF')

def execute():
    # Dešifruoja kanalą ir iškviečia matavimų dešifravimo funkciją
    global stdin, stdout, stderr, channel, flag, current_1_entry, dynamic_changed
    try:
        scope.write(':SYSTem:KEY:PRESs MOFF')
    except NameError:
        info('Oscilloscope is not connected')
    # Apnulina vėliavas
    try:
         if dynamic_changed==1:
            pass
         else: 
            dynamic_changed=0
    except NameError:
        dynamic_changed=0

    try:
        if flag==1:
            pass
        else: 
            flag=0
    except NameError:
        flag=0
    # Kanalų dešifravimas
    if first_channel_variable.get() == 1:
        channel=1
        multiplexer([1,3,2,4],[0,0])
        stdin, stdout, stderr = client.exec_command("ubus call ioman.gpio.dio1 update '{\"value\": \"0\"}'") # Žalias MCU 1
        stdin, stdout, stderr = client.exec_command("ubus call ioman.gpio.dio0 update '{\"value\": \"0\"}'") # Baltas MCU 2
        measurements()
        e_load.write(':SOURce:INPut:STATe OFF')
        root.after(1000)
    
    if second_channel_variable.get() == 1:
        channel=2
        multiplexer([1,3,2,4],[1,0])
        stdin, stdout, stderr = client.exec_command("ubus call ioman.gpio.dio1 update '{\"value\": \"1\"}'") # Žalias MCU 1
        stdin, stdout, stderr = client.exec_command("ubus call ioman.gpio.dio0 update '{\"value\": \"0\"}'") # Baltas MCU 2
        measurements()
        e_load.write(':SOURce:INPut:STATe OFF')
        root.after(1000)
    if third_channel_variable.get() == 1:
        channel=3
        multiplexer([1,3,2,4],[0,1])
        stdin, stdout, stderr = client.exec_command("ubus call ioman.gpio.dio1 update '{\"value\": \"0\"}'") # Žalias MCU 1
        stdin, stdout, stderr = client.exec_command("ubus call ioman.gpio.dio0 update '{\"value\": \"1\"}'") # Baltas MCU 2
        measurements()
        e_load.write(':SOURce:INPut:STATe OFF')
        root.after(1000)
    if fourth_channel_variable.get() == 1:
        channel=4
        multiplexer([1,3,2,4],[1,1])
        stdin, stdout, stderr = client.exec_command("ubus call ioman.gpio.dio1 update '{\"value\": \"1\"}'") # Žalias MCU 1
        stdin, stdout, stderr = client.exec_command("ubus call ioman.gpio.dio0 update '{\"value\": \"1\"}'") # Baltas MCU 2
        measurements()
        e_load.write(':SOURce:INPut:STATe OFF')
        root.after(1000)
    if first_channel_variable.get() == 0 and second_channel_variable.get() == 0 and third_channel_variable.get() == 0 and fourth_channel_variable.get() == 0:
        stdin, stdout, stderr = None
    #00 - 1, 10 -2 , 01 - 3, 11 - 4
    clear_variables()
    info('DONE')
        
# Vartotojo sąsajos pagrindinis laukas
root = Tk()
root.title('Vartotojo sąsaja')
Height=600
Width=700
all_window=Canvas(root, height=500, width=600)
all_window.pack()

# Mažesni laukai, kuriuose atvaizduoti pavadinimai, įvedimo laukeliai
window=Frame(root, bg='#e8fff7')
window.place(relx=0.1, rely=0.05, relwidth=0.8, relheight=0.9)
set_frame=Frame(window)
set_frame.place(relx=0.01, rely=0.1, relwidth=0.5, relheight=0.45)
choose_frame=Frame(window)
choose_frame.place(relx=0.59, rely=0.1, relwidth=0.4, relheight=0.26)#0.27)
Set_current_frame=Frame(window)
Set_current_frame.place(relx=0.01, rely=0.6, relwidth=0.5, relheight=0.39)
current_entry_frame=Frame(window)
current_entry_frame.place(relx=0.59, rely=0.37, relwidth=0.4, relheight=0.24)
Set_dyn_frame=Frame(window)
Set_dyn_frame.place(relx=0.59, rely=0.62, relwidth=0.4, relheight=0.38)
# Teksto langas
text_box = Text(Set_current_frame)
text_box.pack(side=BOTTOM)
# Laukų pavaidinimai
set_lbl=Label(set_frame, text="Set:", bg='#9be8af')
set_lbl.grid(row=0,column=0)
progress_lbl=Label(choose_frame, text="Choose:", bg='#9be8af')
progress_lbl.grid(row=0,column=0, sticky=W)
Title_lbl=Label(window, text="Ripple measurements:", bg='#9be8af')
Title_lbl.pack(side=TOP, anchor=N)
# Pagrindiniai Mygtukai
Start_btn = Button(window, text="Start", command=execute)
Start_btn.pack(side=LEFT)
Cancel_btn = Button(window, text="Cancel", command=cancel)
Cancel_btn.pack(side=LEFT)
Save_btn = Button(window, text="Save as", command=save_file)
Save_btn.pack(side=LEFT)
Change_btn = Button(window, text="Change", command=default_dynamic_parameters)
Change_btn.pack(side=LEFT)
Close_btn = Button(window, text="Close", command=close_connection)
Close_btn.pack(side=LEFT)
# Kintamieji matavimų pasirinkimui
no_load_variable= IntVar()
full_load_variable= IntVar()
dynamic_load_variable= IntVar()
long_term_variable= IntVar()
# Varnelių mygtukai matavimų pasirinkimui
no_load_checkbutton = Checkbutton(choose_frame, text = "No Load", variable = no_load_variable , onvalue = 1, offvalue = 0, state='disabled')
no_load_checkbutton.grid(row=2,column=1, sticky=W)
full_load_checkbutton = Checkbutton(choose_frame, text = "Full Load", variable = full_load_variable , onvalue = 1, offvalue = 0)
full_load_checkbutton.grid(row=3,column=1, sticky=W)
dynamic_load_checkbutton = Checkbutton(choose_frame, text = "Dynamic load", variable = dynamic_load_variable , onvalue = 1, offvalue = 0)
dynamic_load_checkbutton.grid(row=4,column=1, sticky=W)
long_term_checkbutton = Checkbutton(choose_frame, text = "Long Term", variable = long_term_variable , onvalue = 1, offvalue = 0)
long_term_checkbutton.grid(row=5,column=1, sticky=W)


# Kintamieji matinimo šaltinių pasirinkimui
first_channel_variable= IntVar()
second_channel_variable= IntVar()
third_channel_variable= IntVar()
fourth_channel_variable= IntVar()
# Varnelių mygtukai matinimo šaltinių pasirinkimui
Checkbutton(choose_frame, text = "First CH", variable = first_channel_variable , onvalue = 1, offvalue = 0).grid(row=2,column=0, sticky=W)
Checkbutton(choose_frame, text = "Second CH", variable = second_channel_variable , onvalue = 1, offvalue = 0).grid(row=3,column=0, sticky=W)
Checkbutton(choose_frame, text = "Third CH", variable = third_channel_variable , onvalue = 1, offvalue = 0).grid(row=4,column=0, sticky=W)
Checkbutton(choose_frame, text = "Fourth CH", variable = fourth_channel_variable , onvalue = 1, offvalue = 0).grid(row=5,column=0, sticky=W)

#Funkcija kanalų ir matavimo pasirinkimų uždraudimui, kol neprijungti prietaisai
without_current=choose_frame.winfo_children()

for child in without_current[1:]:
            child.configure(state='disabled')

# E apkrovos adreso pasirinkimo pavadinimas, įvedimo langas ir mygtukas
first_address_lbl=Label(set_frame,text='E_load:')
first_address_lbl.grid(row=1,column=0)
first_address_entry = Entry(set_frame, bd =5, width=12)
first_address_entry.grid(row=1,column=1, sticky=W)
first_address_btn = Button(set_frame, text="Set address", command=lambda:[connect(),enable_load_measurements()])
first_address_btn.grid(row=1,column=2, sticky=W)
# Osciloskopo adreso pasirinkimo pavadinimas, įvedimo langas ir mygtukas
second_address_lbl=Label(set_frame,text='Oscilloscope:')
second_address_lbl.grid(row=2,column=0, sticky=W)
second_address_entry = Entry(set_frame, bd =5, width=12)
second_address_entry.grid(row=2,column=1, sticky=W)
second_address_btn = Button(set_frame, text="Set address", command=lambda:[connect(),enable_no_load()])
second_address_btn.grid(row=2,column=2, sticky=W)
# Tinklo sietuvo adreso pasirinkimo pavadinimas, įvedimo langas ir mygtukas
third_address_lbl=Label(set_frame,text='Gateway:')
third_address_lbl.grid(row=3,column=0, sticky=N)
third_address_entry = Entry(set_frame, bd =5, width=12)
third_address_entry.grid(row=3,column=1, sticky=W)
third_address_btn = Button(set_frame, text="Set address", command=connect_trb)
third_address_btn.grid(row=3,column=2, sticky=W)
# Tinklo sietuvo slaptažodžio pasirinkimo pavadinimas, įvedimo langas ir mygtukas
address_password_lbl=Label(set_frame,text='Password:')
address_password_lbl.grid(row=4,column=0, sticky=N)
address_password_entry = Entry(set_frame, bd =5, width=12)
address_password_entry.grid(row=4,column=1, sticky=W)
address_password_btn = Button(set_frame, text="Set password",height = 1,  width = 9,command=lambda:[connect_trb(),enable_load_measurements()])
address_password_btn.grid(row=4,column=2, sticky=W)
# Pradinių adresų ir slaptažodžio reikšmių įvedimas
first_address_entry.insert(END,'192.168.91.20')
second_address_entry.insert(END,'192.168.91.25')
third_address_entry.insert(END,'192.168.91.30')
address_password_entry.insert(END,'Kaunastlt1')

# Srovės nustatymo pavadinimas, įvedimo langas ir mygtukas pirmam kanalui
Current_1_lbl=Label(current_entry_frame,text='Current CH1:')
Current_1_lbl.grid(row=0,column=0)
current_1_entry = Entry(current_entry_frame, bd =5,width=6)
current_1_entry.grid(row=0,column=1)
Set_current_1_btn = Button(current_entry_frame, text="Set current", state='disabled',command=lambda:[set_voltage(),pick_channel(1)])
Set_current_1_btn.grid(row=0,column=2)
# Srovės nustatymo pavadinimas, įvedimo langas ir mygtukas antram kanalui
Current_2_lbl=Label(current_entry_frame,text='Current CH2:')
Current_2_lbl.grid(row=1,column=0)
current_2_entry = Entry(current_entry_frame, bd =5, width=6)
current_2_entry.grid(row=1,column=1)
Set_current_2_btn = Button(current_entry_frame, text="Set current", state='disabled', command=lambda:[set_voltage(),pick_channel(2)])
Set_current_2_btn.grid(row=1,column=2)
# Srovės nustatymo pavadinimas, įvedimo langas ir mygtukas trečiam kanalui
Current_3_lbl=Label(current_entry_frame,text='Current CH3:')
Current_3_lbl.grid(row=2,column=0)
current_3_entry = Entry(current_entry_frame, bd =5, width=6)
current_3_entry.grid(row=2,column=1)
Set_current_3_btn = Button(current_entry_frame, text="Set current", state='disabled', command=lambda:[set_voltage(),pick_channel(3)])
Set_current_3_btn.grid(row=2,column=2)
# Srovės nustatymo pavadinimas, įvedimo langas ir mygtukas ketvirtam kanalui
Current_4_lbl=Label(current_entry_frame,text='Current CH4:')
Current_4_lbl.grid(row=3,column=0)
current_4_entry = Entry(current_entry_frame, bd =5, width=6)
current_4_entry.grid(row=3,column=1)
Set_current_4_btn = Button(current_entry_frame, text="Set current", state='disabled', command=lambda:[set_voltage(),pick_channel(4)])
Set_current_4_btn.grid(row=3,column=2)

# Įtampos nustatymo pavadinimas, įvedimo langas ir mygtukas 
Voltage_lbl=Label(set_frame,text='Voltage:')
Voltage_lbl.grid(row=6,column=0)
voltage_entry = Entry(set_frame, bd =5, width=12)
voltage_entry.grid(row=6,column=1)
Set_voltage_btn = Button(set_frame, text="Set voltage", command=set_voltage)
Set_voltage_btn.grid(row=6,column=2)
# Įšėjimo mygtukas
Button(root, text="Quit", command=root.destroy).pack()
mainloop()